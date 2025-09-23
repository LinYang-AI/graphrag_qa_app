import asyncio
import hashlib
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
import logging
import re

import aiofiles
from docx import Document as DocxDocument
from openpyxl import load_workbook
from bs4 import BeautifulSoup
import markdown
import PyPDF2
from PIL import Image

logger = logging.getLogger(__name__)


class DocumentMetadata:
    """Structured document metadata"""
    def __init__(self):
        self.title: Optional[str] = None
        self.author: Optional[str] = None
        self.created_date: Optional[datetime] = None
        self.modified_date: Optional[datetime] = None
        self.file_size: int = 0
        self.file_type: Optional[str] = None
        self.mime_type: Optional[str] = None
        self.page_count: int = 0
        self.word_count: int = 0
        self.language: str = "en"
        self.tenant_id: str = "default"
        self.document_hash: Optional[str] = None
        self.version: int = 1
        self.tags: List[str] = []

    def to_dict(self) -> Dict[str, Any]:
        """Convert metadata to dictionary for storage"""
        return {
            "title": self.title,
            "author": self.author,
            "created_date": self.created_date.isoformat() if self.created_date else None,
            "modified_date": self.modified_date.isoformat() if self.modified_date else None,
            "file_size": self.file_size,
            "file_type": self.file_type,
            "mime_type": self.mime_type,
            "page_count": self.page_count,
            "word_count": self.word_count,
            "language": self.language,
            "tenant_id": self.tenant_id,
            "document_hash": self.document_hash,
            "version": self.version,
            "tags": self.tags
        }


class DocumentChunk:
    """Represents a processed document chunk with metadata"""
    def __init__(self, content: str, chunk_index: int, metadata: DocumentMetadata):
        self.content = content
        self.chunk_index = chunk_index
        self.metadata = metadata
        self.chunk_id = f"{metadata.document_hash}_{chunk_index}"
        self.word_count = len(content.split())
        self.char_count = len(content)
        
        # Extract section information if available
        self.section_title = self._extract_section_title(content)
        self.chunk_type = self._determine_chunk_type(content)

    def _extract_section_title(self, content: str) -> Optional[str]:
        """Try to extract section title from chunk content"""
        lines = content.strip().split('\n')
        first_line = lines[0].strip()
        
        # Check if first line looks like a title/header
        if (len(first_line) < 100 and 
            (first_line.isupper() or 
             re.match(r'^#{1,6}\s+', first_line) or  # Markdown headers
             re.match(r'^\d+\.?\s+[A-Z]', first_line))):  # Numbered sections
            return first_line
        return None

    def _determine_chunk_type(self, content: str) -> str:
        """Determine the type of content in this chunk"""
        content_lower = content.lower()
        
        if any(keyword in content_lower for keyword in ['table', 'figure', 'chart', 'graph']):
            return "data_visualization"
        elif any(keyword in content_lower for keyword in ['conclusion', 'summary', 'abstract']):
            return "summary"
        elif re.search(r'\b\d{4}-\d{2}-\d{2}\b', content) or 'date' in content_lower:
            return "temporal"
        elif len(re.findall(r'\$\d+', content)) > 2:
            return "financial"
        else:
            return "general"

    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary for storage"""
        return {
            "content": self.content,
            "chunk_id": self.chunk_id,
            "chunk_index": self.chunk_index,
            "word_count": self.word_count,
            "char_count": self.char_count,
            "section_title": self.section_title,
            "chunk_type": self.chunk_type,
            "metadata": self.metadata.to_dict()
        }


class AdvancedDocumentProcessor:
    """Advanced document processing with multi-format support and rich metadata extraction"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.xlsx': self._process_xlsx,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html
        }
        self.chunk_strategies = {
            'paragraph': self._chunk_by_paragraphs,
            'sentence': self._chunk_by_sentences,
            'fixed': self._chunk_by_fixed_size,
            'semantic': self._chunk_by_semantic_boundaries
        }

    async def process_document(self, file_path: str, tenant_id: str = "default", 
                             chunk_strategy: str = "paragraph") -> Tuple[DocumentMetadata, List[DocumentChunk]]:
        """Process a document and return metadata + chunks"""
        try:
            logger.info(f"📄 Processing document: {Path(file_path).name}")
            
            # Extract basic file info
            metadata = await self._extract_basic_metadata(file_path, tenant_id)
            
            # Extract content based on file type
            content = await self._extract_content(file_path, metadata)
            
            if not content or len(content.strip()) < 50:
                logger.warning(f"Insufficient content extracted from {file_path}")
                return metadata, []

            # Update metadata with content analysis
            metadata.word_count = len(content.split())
            metadata.document_hash = self._generate_document_hash(content)

            # Create chunks using specified strategy
            chunks = await self._create_chunks(content, metadata, chunk_strategy)
            
            logger.info(f"✅ Processed {Path(file_path).name}: {len(chunks)} chunks, {metadata.word_count} words")
            return metadata, chunks

        except Exception as e:
            logger.error(f"❌ Error processing {file_path}: {e}")
            raise

    async def _extract_basic_metadata(self, file_path: str, tenant_id: str) -> DocumentMetadata:
        """Extract basic file system metadata"""
        path = Path(file_path)
        stat = path.stat()
        
        metadata = DocumentMetadata()
        metadata.title = path.stem
        metadata.file_type = path.suffix.lower()
        metadata.mime_type = mimetypes.guess_type(file_path)[0]
        metadata.file_size = stat.st_size
        metadata.created_date = datetime.fromtimestamp(stat.st_ctime)
        metadata.modified_date = datetime.fromtimestamp(stat.st_mtime)
        metadata.tenant_id = tenant_id
        
        return metadata

    async def _extract_content(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Extract content based on file type"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext in self.supported_formats:
            processor = self.supported_formats[file_ext]
            content = await processor(file_path, metadata)
            return content
        else:
            logger.warning(f"Unsupported file format: {file_ext}")
            return ""

    async def _process_pdf(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Extract text from PDF with metadata"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata.page_count = len(pdf_reader.pages)
                
                # Extract PDF metadata
                if pdf_reader.metadata:
                    pdf_info = pdf_reader.metadata
                    metadata.title = pdf_info.get('/Title', metadata.title)
                    metadata.author = pdf_info.get('/Author', None)
                    
                    # Handle creation date
                    if '/CreationDate' in pdf_info:
                        try:
                            # PDF dates are in format: D:YYYYMMDDHHmmSSOHH'mm'
                            date_str = str(pdf_info['/CreationDate']).replace('D:', '')[:14]
                            metadata.created_date = datetime.strptime(date_str, '%Y%m%d%H%M%S')
                        except:
                            pass  # Keep file system date
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting page {page_num + 1}: {e}")
                
                return '\n\n'.join(text_content)
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return ""

    async def _process_docx(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Extract text from DOCX with metadata"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract document properties
            core_props = doc.core_properties
            metadata.title = core_props.title or metadata.title
            metadata.author = core_props.author
            metadata.created_date = core_props.created or metadata.created_date
            metadata.modified_date = core_props.modified or metadata.modified_date
            
            # Extract text content preserving structure
            content_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Check if it's a heading
                    if paragraph.style.name.startswith('Heading'):
                        content_parts.append(f"\n## {paragraph.text}\n")
                    else:
                        content_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                
                if table_text:
                    content_parts.append("\n--- Table ---\n" + "\n".join(table_text) + "\n")
            
            return '\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""

    async def _process_xlsx(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Extract text from Excel file"""
        try:
            workbook = load_workbook(file_path, read_only=True)
            
            # Extract workbook properties
            props = workbook.properties
            metadata.title = props.title or metadata.title
            metadata.author = props.creator
            metadata.created_date = props.created or metadata.created_date
            metadata.modified_date = props.modified or metadata.modified_date
            
            content_parts = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content_parts.append(f"\n=== Sheet: {sheet_name} ===\n")
                
                # Extract data with headers
                rows_data = []
                for row_num, row in enumerate(sheet.iter_rows(values_only=True), 1):
                    if any(cell is not None for cell in row):
                        row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                        rows_data.append(row_text)
                        
                        # Limit rows to avoid huge content
                        if row_num > 100:
                            rows_data.append("... (truncated)")
                            break
                
                content_parts.extend(rows_data)
            
            workbook.close()
            return '\n'.join(content_parts)
            
        except Exception as e:
            logger.error(f"Error processing XLSX {file_path}: {e}")
            return ""

    async def _process_text(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Process plain text file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
                return content.strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    content = await file.read()
                    return content.strip()
            except Exception as e:
                logger.error(f"Error reading text file {file_path}: {e}")
                return ""

    async def _process_markdown(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Process Markdown file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                md_content = await file.read()
                
                # Convert markdown to HTML then extract text, preserving structure
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                
                # Convert HTML structure to plain text with markers
                content_parts = []
                for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li']):
                    if element.name.startswith('h'):
                        level = element.name[1]
                        content_parts.append(f"\n{'#' * int(level)} {element.get_text().strip()}\n")
                    elif element.name == 'li':
                        content_parts.append(f"• {element.get_text().strip()}")
                    else:
                        text = element.get_text().strip()
                        if text:
                            content_parts.append(text)
                
                return '\n'.join(content_parts)
                
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}")
            return ""

    async def _process_html(self, file_path: str, metadata: DocumentMetadata) -> str:
        """Process HTML file"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                html_content = await file.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Extract metadata from HTML
                title_tag = soup.find('title')
                if title_tag:
                    metadata.title = title_tag.get_text().strip()
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Extract text content
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                return text
                
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            return ""

    def _generate_document_hash(self, content: str) -> str:
        """Generate a hash for document deduplication"""
        return hashlib.sha256(content.encode('utf-8')).hexdigest()[:16]

    async def _create_chunks(self, content: str, metadata: DocumentMetadata, 
                           strategy: str = "paragraph") -> List[DocumentChunk]:
        """Create chunks using specified strategy"""
        if strategy not in self.chunk_strategies:
            logger.warning(f"Unknown chunk strategy: {strategy}, using 'paragraph'")
            strategy = "paragraph"
        
        chunker = self.chunk_strategies[strategy]
        chunk_texts = await chunker(content)
        
        # Create DocumentChunk objects
        chunks = []
        for i, chunk_text in enumerate(chunk_texts):
            if len(chunk_text.strip()) > 50:  # Filter out very short chunks
                chunk = DocumentChunk(chunk_text, i, metadata)
                chunks.append(chunk)
        
        return chunks

    async def _chunk_by_paragraphs(self, content: str, max_chunk_size: int = 800) -> List[str]:
        """Smart paragraph-based chunking that respects document structure"""
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_size = 0
        
        for paragraph in paragraphs:
            paragraph_size = len(paragraph)
            
            # If adding this paragraph would exceed max size, finalize current chunk
            if current_size + paragraph_size > max_chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = [paragraph]
                current_size = paragraph_size
            else:
                current_chunk.append(paragraph)
                current_size += paragraph_size + 2  # +2 for \n\n
        
        # Add the last chunk
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        return chunks

    async def _chunk_by_sentences(self, content: str, max_chunk_size: int = 600) -> List[str]:
        """Sentence-based chunking for better semantic coherence"""
        # Simple sentence splitting (can be enhanced with spaCy)
        sentences = re.split(r'(?<=[.!?])\s+', content)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_size = len(sentence)
            
            if current_size + sentence_size > max_chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_size = sentence_size
            else:
                current_chunk.append(sentence)
                current_size += sentence_size + 1  # +1 for space
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    async def _chunk_by_fixed_size(self, content: str, chunk_size: int = 500, overlap: int = 50) -> List[str]:
        """Fixed-size chunking with overlap (fallback method)"""
        words = content.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk)
        
        return chunks

    async def _chunk_by_semantic_boundaries(self, content: str, max_chunk_size: int = 700) -> List[str]:
        """Semantic boundary chunking (looks for section headers, lists, etc.)"""
        # Split by common semantic boundaries
        sections = re.split(r'\n(?=#{1,6}\s|^\d+\.\s|^[A-Z][^a-z]*$)', content, flags=re.MULTILINE)
        chunks = []
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            if len(section) <= max_chunk_size:
                chunks.append(section)
            else:
                # If section is too large, fall back to paragraph chunking
                sub_chunks = await self._chunk_by_paragraphs(section, max_chunk_size)
                chunks.extend(sub_chunks)
        
        return chunks

    async def process_batch(self, file_paths: List[str], tenant_id: str = "default",
                          chunk_strategy: str = "paragraph") -> Dict[str, Any]:
        """Process multiple documents in batch"""
        results = {
            "processed": 0,
            "failed": 0,
            "total_chunks": 0,
            "documents": [],
            "errors": []
        }
        
        for file_path in file_paths:
            try:
                metadata, chunks = await self.process_document(file_path, tenant_id, chunk_strategy)
                results["processed"] += 1
                results["total_chunks"] += len(chunks)
                results["documents"].append({
                    "file_path": file_path,
                    "metadata": metadata.to_dict(),
                    "chunk_count": len(chunks)
                })
            except Exception as e:
                results["failed"] += 1
                results["errors"].append({
                    "file_path": file_path,
                    "error": str(e)
                })
                logger.error(f"Failed to process {file_path}: {e}")
        
        return results
